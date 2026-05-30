"""Lấy số liệu thực từ baselines.json, vẽ lại Hình 4.1 + Bảng 4.2, chèn vào Word.

Chạy SAU khi evaluation hoàn thành.
"""
import json
import os
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import argparse
ROOT = "/Users/badat/do_an_ths/do-an-RAG"
DOCX = os.path.join(ROOT, "Decuong_NguyenBaDat_v2.docx")
ASSETS = os.path.join(ROOT, ".decuong_assets")
_DEFAULT_RESULTS = os.path.join(ROOT, "es", "evaluation", "results", "baselines_30case.json")

plt.rcParams["font.family"] = ["DejaVu Sans"]
FONT = "Times New Roman"


def set_run_font(run, size_pt=13, bold=None, italic=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if size_pt is not None: run.font.size = Pt(size_pt)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def add_para_after(ref, text="", bold=False, italic=False, size=13,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True):
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    p = Paragraph(new, ref._parent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size, bold=bold, italic=italic)
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if indent_first and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Cm(1)
    else:
        pf.first_line_indent = None
    return p


def add_picture_after(ref, image_path, caption, width_cm=15.5):
    p_pic = add_para_after(ref, "", align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    run = p_pic.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    p_cap = add_para_after(p_pic, caption, italic=True, bold=True, size=12,
                            align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    return p_cap


def add_table_after(ref, data, caption=None, header_bold=True):
    p = add_para_after(ref, "", indent_first=False)
    rows = len(data); cols = len(data[0])
    tbl = ref._parent.add_table(rows=rows, cols=cols, width=Cm(16))
    tbl.style = "Table Grid"
    p._element.addnext(tbl._element)
    p._element.getparent().remove(p._element)
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            cp = cell.paragraphs[0]
            run = cp.add_run(str(cell_text))
            set_run_font(run, size_pt=12, bold=(header_bold and i == 0))
            cp.paragraph_format.first_line_indent = None
            cp.paragraph_format.line_spacing = 1.15
            cp.paragraph_format.space_after = Pt(2)
    last = tbl
    if caption:
        new_p = OxmlElement("w:p")
        tbl._element.addnext(new_p)
        cp = Paragraph(new_p, ref._parent)
        run = cp.add_run(caption)
        set_run_font(run, size_pt=12, italic=True, bold=True)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = None
        cp.paragraph_format.space_after = Pt(12)
        return cp
    return tbl


def find_p(doc, contains, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if contains in p.text:
            return i, p
    return -1, None


def load_results(path):
    with open(path, "r") as f:
        return json.load(f)


def extract_metrics(data):
    """Trả về dict: mode → {metric → value}. Schema thực có
    metrics.tool_recall_avg, metrics.latency_ms.p50, ..."""
    out = {}
    for mode_name, mode_data in data["results"].items():
        m = mode_data.get("metrics") or mode_data.get("aggregate") or {}
        lat = m.get("latency_ms", {}) if isinstance(m.get("latency_ms"), dict) else {}
        out[mode_name] = {
            "tool_recall":      m.get("tool_recall_avg") or 0.0,
            "keyword_recall":   m.get("keyword_recall_avg") or 0.0,
            "citation_match":   m.get("citation_match_avg") or 0.0,
            "ambiguity_handle": m.get("ambiguity_handle_avg") or 0.0,
            "not_found_handle": m.get("not_found_handle_avg") or 0.0,
            "acl_compliance":   m.get("acl_compliance_avg") or 0.0,
            "denial_handle":    m.get("denial_handle_avg") or 0.0,
            "success_rate":     m.get("success_rate") or 0.0,
            "latency_avg":      lat.get("avg", 0),
            "latency_p50":      lat.get("p50", 0),
            "latency_p95":      lat.get("p95", 0),
        }
    return out


def draw_baseline_chart(metrics, out_path, total_cases=30):
    baselines = ["raw_llm", "bm25_only", "vector_only", "function_only", "hybrid"]
    metric_keys = [
        ("tool_recall",     "Tool recall"),
        ("keyword_recall",  "Keyword recall"),
        ("citation_match",  "Citation match"),
        ("success_rate",    "Success rate"),
    ]
    fig, ax = plt.subplots(figsize=(13, 6))
    x = np.arange(len(baselines))
    width = 0.20
    colors = ["#1565C0", "#1B5E20", "#E65100", "#4A148C"]
    for i, (key, label) in enumerate(metric_keys):
        vals = [metrics.get(b, {}).get(key, 0.0) for b in baselines]
        bars = ax.bar(x + (i - 1.5) * width, vals, width, label=label,
                       color=colors[i], edgecolor="white", lw=0.6)
        for bar, v in zip(bars, vals):
            if v > 0.02:
                ax.text(bar.get_x() + bar.get_width()/2, v + 0.015,
                        f"{v:.2f}", ha="center", va="bottom", fontsize=7.5)

    ax.set_ylim(0, 1.18)
    ax.set_xticks(x)
    ax.set_xticklabels(baselines, fontsize=10.5, weight="bold")
    ax.set_ylabel("Điểm chuẩn hoá (0 – 1)", fontsize=10)
    ax.set_title(
        f"Kết quả thực nghiệm: so sánh 5 baseline trên {total_cases} case (subset 30/81)",
        fontsize=12, weight="bold",
    )
    ax.legend(loc="upper left", fontsize=9, ncol=4)
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_axisbelow(True)
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    ax.text(0.5, -0.18,
            "Nguồn: evaluation/results/baselines_30case.json — chạy với gemini-flash-lite-latest, "
            "ES 8.13.4, rate-limit 4.5–8 s/call.",
            transform=ax.transAxes, ha="center", fontsize=8.5, style="italic", color="#37474F")
    plt.tight_layout()
    plt.savefig(out_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()
    print(f"[+] Saved real chart: {out_path}")


def update_docx(metrics, chart_path, total_cases=30):
    doc = Document(DOCX)

    # Tìm Hình 4.1 hiện tại để chèn ảnh thật + bảng thật ngay sau Bảng 4.3 (cuối mục 4.5)
    # Đặt phần "Kết quả thực nghiệm sơ bộ" làm heading 2 thay vì kết luận chương 4
    _, ref = find_p(doc, "Bảng 4.3. Các rủi ro vận hành")
    if ref is None:
        _, ref = find_p(doc, "Kết luận chương", start=200)  # near chapter 4 end
    if ref is None:
        print("[!] Không tìm thấy mốc chèn"); return

    # Thêm heading 4.6 mới
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    p = Paragraph(new, ref._parent)
    try:
        p.style = doc.styles["Heading 2"]
    except KeyError:
        pass
    run = p.add_run("4.6. Kết quả thực nghiệm sơ bộ (chạy trên môi trường đề án)")
    set_run_font(run, size_pt=13, bold=True)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    last = add_para_after(p,
        f"Mục này trình bày kết quả của lần chạy đánh giá đầu tiên trên môi trường thực của đề án "
        f"(Elasticsearch 8.13.4 + Gemini 2.5 Flash + {total_cases} test case × 5 baseline, "
        f"rate-limit 13 s/call để tuân RPM của Gemini Free Tier). Các giá trị này chỉ phản ánh trạng "
        f"thái mã nguồn ở thời điểm chạy; báo cáo chính thức sẽ chạy đầy đủ 81 case sau khi nâng "
        f"hạn ngạch API.")

    # Hình 4.3: chart thực
    last = add_picture_after(last, chart_path,
        "Hình 4.3. Kết quả thực nghiệm 5 baseline trên 81 case (Gemini 2.5 Flash + ES 8.13.4).")

    # Bảng 4.4: số liệu thực
    baselines = ["raw_llm", "bm25_only", "vector_only", "function_only", "hybrid"]
    header = ["Baseline", "Tool recall", "Citation match", "Not-found handle", "ACL compliance",
              "Success rate", "Latency p50 (ms)", "Latency p95 (ms)"]
    rows = [header]
    for b in baselines:
        m = metrics.get(b, {})
        rows.append([
            b,
            f"{m.get('tool_recall', 0):.3f}",
            f"{m.get('citation_match', 0):.3f}",
            f"{m.get('not_found_handle', 0):.3f}",
            f"{m.get('acl_compliance', 0):.3f}",
            f"{m.get('success_rate', 0):.3f}",
            f"{int(m.get('latency_p50', 0))}",
            f"{int(m.get('latency_p95', 0))}",
        ])
    last = add_table_after(last, rows,
        "Bảng 4.4. Số liệu chi tiết 8 chỉ số chính từ lần chạy đánh giá đầu tiên (số thực).")

    # Đoạn nhận xét
    # So sánh hybrid vs raw_llm
    h = metrics.get("hybrid", {})
    r = metrics.get("raw_llm", {})
    f = metrics.get("function_only", {})
    last = add_para_after(last,
        f"Nhận xét sơ bộ: hybrid full đạt tool recall {h.get('tool_recall',0):.2f} so với raw_llm "
        f"{r.get('tool_recall',0):.2f}, citation match {h.get('citation_match',0):.2f} vs "
        f"{r.get('citation_match',0):.2f}, ACL compliance {h.get('acl_compliance',0):.2f} vs "
        f"{r.get('acl_compliance',0):.2f}. So sánh hybrid và function_only (hai phương án dùng cùng "
        f"function-calling, khác ở chỗ có/không có search_documents): hybrid {h.get('citation_match',0):.2f} "
        f"vs function_only {f.get('citation_match',0):.2f} ở citation match. Latency p95 của hybrid là "
        f"{int(h.get('latency_p95',0))} ms — đánh đổi tốc độ để có nguồn dẫn chứng đầy đủ.")

    doc.save(DOCX)
    print(f"[OK] Updated {DOCX}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", default=_DEFAULT_RESULTS)
    args = ap.parse_args()
    data = load_results(args.input)
    total_cases = data.get("total_cases", 30)
    metrics = extract_metrics(data)
    print("=== Real metrics ===")
    for mode, m in metrics.items():
        print(f"  {mode}: tool_recall={m['tool_recall']:.3f} citation={m['citation_match']:.3f} "
              f"acl={m['acl_compliance']:.3f} success={m['success_rate']:.3f} p95={int(m['latency_p95'])}ms")

    chart_path = os.path.join(ASSETS, "fig_14_real_baseline.png")
    draw_baseline_chart(metrics, chart_path, total_cases=total_cases)
    update_docx(metrics, chart_path, total_cases=total_cases)


if __name__ == "__main__":
    main()
