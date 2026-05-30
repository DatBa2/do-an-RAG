"""Bổ sung phần Elasticsearch:
1. Chèn Hình 1.8 (ES one-store) vào mục 1.4
2. Chèn đoạn lý do chọn ES vào mục 1.4 (sau bảng so sánh vector DB)
3. Thêm Phụ lục A: mapping JSON
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = "/Users/badat/do_an_ths/do-an-RAG"
SRC = os.path.join(ROOT, "Decuong_NguyenBaDat_v2.docx")
ASSETS = os.path.join(ROOT, ".decuong_assets")
FONT = "Times New Roman"


def set_run_font(run, size_pt=13, bold=None, italic=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if size_pt is not None: run.font.size = Pt(size_pt)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def add_para_after(ref, text="", bold=False, italic=False, size=13,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
                   mono=False, line_spacing=1.5):
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    p = Paragraph(new, ref._parent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size, bold=bold, italic=italic)
        if mono:
            run.font.name = "Consolas"
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rFonts.set(qn(attr), "Consolas")
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(4) if mono else Pt(6)
    pf.space_before = Pt(0)
    if indent_first and align == WD_ALIGN_PARAGRAPH.JUSTIFY and not mono:
        pf.first_line_indent = Cm(1)
    else:
        pf.first_line_indent = None
    return p


def add_picture_after(ref, image_path, caption, width_cm=15.5):
    p_pic = add_para_after(ref, "", align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    run = p_pic.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    p_cap = add_para_after(p_pic, caption, italic=True, bold=True, size=12,
                            align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    return p_cap


def add_heading_after(ref, text, level=2):
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    p = Paragraph(new, ref._parent)
    p.style = ref._parent.part.document.styles[f"Heading {level}"]
    run = p.add_run(text)
    set_run_font(run, size_pt=14 if level == 1 else 13, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def find_p(doc, contains, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if contains in p.text:
            return i, p
    return -1, None


# Mapping JSON content
ES_MAPPING_HS_RECORDS = """\
PUT /hs_records
{
  "settings": {
    "number_of_shards": 1,
    "number_of_replicas": 1,
    "analysis": {
      "analyzer": {
        "vn_text": {
          "type": "custom",
          "tokenizer": "standard",
          "filter": ["lowercase", "asciifolding", "vn_synonyms", "stop"]
        }
      },
      "filter": {
        "vn_synonyms": {
          "type": "synonym_graph",
          "synonyms": [
            "HK1, học kỳ 1, học kì 1, hk1",
            "HK2, học kỳ 2, học kì 2, hk2",
            "TB, ĐTB, GPA, trung bình",
            "GVCN, giáo viên chủ nhiệm, chủ nhiệm",
            "TX, thường xuyên",
            "GK, giữa kỳ",
            "CK, cuối kỳ",
            "TK, tổng kết",
            "HS, học sinh"
          ]
        }
      }
    }
  },
  "mappings": {
    "properties": {
      "doc_type":         { "type": "keyword" },
      "student_id":       { "type": "keyword" },
      "full_name":        {
        "type": "text", "analyzer": "vn_text",
        "fields": { "raw": { "type": "keyword" } }
      },
      "class_name":       {
        "type": "text", "analyzer": "vn_text",
        "fields": { "raw": { "type": "keyword" } }
      },
      "year":             { "type": "keyword" },
      "semester":         { "type": "integer" },
      "subject":          {
        "type": "text", "analyzer": "vn_text",
        "fields": { "raw": { "type": "keyword" } }
      },
      "scores": {
        "properties": {
          "TX": { "type": "float" },
          "GK": { "type": "float" },
          "CK": { "type": "float" },
          "TK": { "type": "float" }
        }
      },
      "overall_gpa":      { "type": "float" },
      "conduct":          { "type": "keyword" },
      "academic":         { "type": "keyword" },
      "promotion":        { "type": "keyword" },
      "attendance": {
        "properties": {
          "phep":        { "type": "integer" },
          "khong_phep":  { "type": "integer" },
          "bo_tiet":     { "type": "integer" }
        }
      },
      "homeroom_comment": { "type": "text", "analyzer": "vn_text" },
      "raw_path":         { "type": "keyword", "index": false }
    }
  }
}"""

ES_MAPPING_INTERNAL_DOCS = """\
PUT /internal_docs
{
  "settings": {
    "number_of_shards": 1,
    "number_of_replicas": 1,
    "analysis": {
      "analyzer": {
        "vn_text": {
          "type": "custom",
          "tokenizer": "standard",
          "filter": ["lowercase", "asciifolding", "vn_synonyms", "stop"]
        }
      },
      "filter": {
        "vn_synonyms": {
          "type": "synonym_graph",
          "synonyms_path": "analysis/synonyms_vn.txt"
        }
      }
    }
  },
  "mappings": {
    "properties": {
      "doc_id":         { "type": "keyword" },
      "chunk_id":       { "type": "keyword" },
      "title":          { "type": "text", "analyzer": "vn_text" },
      "section":        { "type": "text", "analyzer": "vn_text" },
      "source_path":    { "type": "keyword" },
      "content":        { "type": "text", "analyzer": "vn_text" },
      "embedding": {
        "type": "dense_vector",
        "dims": 768,
        "index": true,
        "similarity": "cosine",
        "index_options": { "type": "int8_hnsw", "m": 16, "ef_construction": 100 }
      },
      "content_tokens": { "type": "integer" },
      "allowed_roles":  { "type": "keyword" },
      "indexed_at":     { "type": "date" }
    }
  }
}"""

ES_HYBRID_RETRIEVER = """\
POST /internal_docs/_search
{
  "retriever": {
    "rrf": {
      "rank_constant": 60,
      "rank_window_size": 50,
      "retrievers": [
        {
          "standard": {
            "query": {
              "bool": {
                "must":   { "match": { "content": { "query": "<câu hỏi>" } } },
                "filter": { "terms": { "allowed_roles": ["<role>"] } }
              }
            }
          }
        },
        {
          "knn": {
            "field": "embedding",
            "query_vector": [/* 768 floats từ retrieval_query */],
            "k": 50,
            "num_candidates": 200,
            "filter": { "terms": { "allowed_roles": ["<role>"] } }
          }
        }
      ]
    }
  },
  "size": 5,
  "_source": ["doc_id", "title", "section", "source_path", "content"]
}"""


def main():
    doc = Document(SRC)

    # ============================================================
    # 1. Chèn Hình ES one-store + đoạn lý do chọn ES sau bảng 1.3
    # ============================================================
    _, ref = find_p(doc, "Bảng 1.3. So sánh các giải pháp lưu trữ vector")
    if ref is not None:
        last = add_picture_after(ref,
            os.path.join(ASSETS, "fig_13_es_onestore.png"),
            "Hình 1.8. Elasticsearch trong vai trò 'one-store' cho cả BM25 và KNN trong đề án.")

        # 3 đoạn lý do chọn ES
        last = add_para_after(last,
            "Tại sao đề án chọn Elasticsearch dù các vector DB chuyên dụng (Qdrant, Milvus, Pinecone) đang phát triển nhanh?",
            bold=True, indent_first=False)

        last = add_para_after(last,
            "Thứ nhất, mô hình one-store hai retrieval. Hệ thống của đề án phục vụ đồng thời hai luồng truy xuất: BM25 cho "
            "câu hỏi giàu từ khoá - tên riêng, mã môn, số liệu - và KNN cho câu hỏi diễn đạt ngữ nghĩa. Elasticsearch 8.x là "
            "lựa chọn hiếm hoi cung cấp cả hai trong cùng một dịch vụ; các vector DB chuyên dụng như Qdrant hay Milvus phải "
            "ghép thêm một engine BM25 riêng (Elasticsearch, OpenSearch, hoặc PostgreSQL tsvector) - làm tăng độ phức tạp "
            "vận hành. Bản thân ES từ phiên bản 8.8 đã có retriever API hỗ trợ Reciprocal Rank Fusion native: một call HTTP "
            "duy nhất sinh ra danh sách hợp nhất, không cần xử lý phía client.")

        last = add_para_after(last,
            "Thứ hai, đồng bộ ACL ở một chỗ duy nhất. Trong kiến trúc nhiều dịch vụ, việc đảm bảo bảng allowed_roles được "
            "đồng bộ chính xác giữa hai store (text store và vector store) là một bài toán không tầm thường - sai sót có thể "
            "dẫn đến rò rỉ tài liệu vượt quyền theo nhánh vector mà nhánh text đã chặn. Elasticsearch giữ allowed_roles ngay "
            "trong cùng một document JSON cho cả trường content (phục vụ BM25) và trường embedding (phục vụ KNN). Filter "
            "trong câu truy vấn áp lên cả hai nhánh đồng thời, tránh được hoàn toàn nguy cơ phân kỳ ACL - phù hợp với khuyến "
            "nghị 'Vector and Embedding Weaknesses' của OWASP LLM Top 10 - 2025 (LLM08).")

        last = add_para_after(last,
            "Thứ ba, hệ sinh thái sản xuất chín muồi. Elastic Stack đã được triển khai rộng rãi trong các tổ chức Việt Nam "
            "phục vụ log analytics, SIEM, full-text search; đội ngũ DevOps quen với ES dễ tiếp quản hệ thống RAG mà không "
            "cần học công cụ mới. Các vector DB chuyên dụng có hiệu năng tốt hơn ở quy mô hàng tỷ vector, nhưng quy mô của "
            "đề án (~10^3 - 10^5 vector) chưa chạm ngưỡng đó. Bản thân kiến trúc của đề án vẫn cho phép swap sang Qdrant "
            "hay Milvus trong tương lai nhờ module doc_search.py đóng gói ba mode (bm25, vector, hybrid) sau một interface "
            "duy nhất - chỉ cần thay phần triển khai bên trong.")

        last = add_para_after(last,
            "Thứ tư, chi phí thấp - tự host. ES có thể chạy trên một node Docker đơn giản cho giai đoạn POC, mở rộng dần "
            "sang cluster nhiều node khi tăng tải. Pinecone (managed cloud) cho UX tốt nhưng chi phí trên 1 triệu vector + "
            "request/tháng cao hơn đáng kể so với một máy chủ tự host. Đối với một trường THCS, chi phí vận hành là yếu tố "
            "then chốt; đề án ưu tiên giải pháp on-premise.")
        print("[+] Đã chèn hình ES one-store + 4 đoạn lý do chọn ES")

    # ============================================================
    # 2. Phụ lục A: mapping JSON
    # ============================================================
    # Tìm vị trí cuối Phần V (sau bảng kế hoạch) hoặc trước heading cuối
    _, ref = find_p(doc, "DUYỆT CỦA TRƯỞNG TIỂU BAN")
    if ref is None:
        _, ref = find_p(doc, "PGS.TS.")
    if ref is None:
        # fallback: cuối tài liệu
        ref = doc.paragraphs[-1]

    # Insert page break before appendix
    from docx.enum.text import WD_BREAK
    pb_p = add_para_after(ref, "", indent_first=False)
    pb_run = pb_p.add_run()
    pb_run.add_break(WD_BREAK.PAGE)

    last = add_heading_after(pb_p, "PHỤ LỤC A. MAPPING ELASTICSEARCH VÀ TRUY VẤN HYBRID", level=1)

    last = add_para_after(last,
        "Phụ lục này trình bày mapping JSON đầy đủ của hai index Elasticsearch trong đề án cùng câu truy vấn hybrid mẫu, "
        "phục vụ tham chiếu kỹ thuật và tái lập thực nghiệm.")

    # A.1 hs_records
    last = add_heading_after(last, "A.1. Mapping index hs_records (dữ liệu cấu trúc)", level=2)
    last = add_para_after(last,
        "Index hs_records lưu hồ sơ học sinh và bảng điểm dạng JSON. Mỗi document có doc_type = 'student' hoặc 'mark'. "
        "Analyzer vn_text áp dụng asciifolding (loại dấu) và bộ từ đồng nghĩa giáo dục.")
    # Chèn từng dòng JSON (giữ nguyên xuống dòng) bằng nhiều paragraph
    for line in ES_MAPPING_HS_RECORDS.splitlines():
        last = add_para_after(last, line or " ", mono=True, size=10,
                              align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False,
                              line_spacing=1.1)

    # A.2 internal_docs
    last = add_heading_after(last, "A.2. Mapping index internal_docs (tài liệu phi cấu trúc)", level=2)
    last = add_para_after(last,
        "Index internal_docs lưu các chunk văn bản kèm embedding 768 chiều và trường allowed_roles cho ACL. Trường "
        "embedding khai báo index_options dạng int8_hnsw để nén 4 lần dung lượng đĩa với mất mát chính xác không đáng kể.")
    for line in ES_MAPPING_INTERNAL_DOCS.splitlines():
        last = add_para_after(last, line or " ", mono=True, size=10,
                              align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False,
                              line_spacing=1.1)

    # A.3 Hybrid query
    last = add_heading_after(last, "A.3. Câu truy vấn Hybrid retriever (RRF native từ ES 8.8)", level=2)
    last = add_para_after(last,
        "Câu truy vấn dưới đây phối hợp BM25 và KNN trên cùng một index, lọc theo allowed_roles ở cả hai nhánh, sau đó "
        "hợp nhất bằng Reciprocal Rank Fusion với rank_constant = 60 - giá trị mặc định khuyến nghị từ Cormack et al. 2009.")
    for line in ES_HYBRID_RETRIEVER.splitlines():
        last = add_para_after(last, line or " ", mono=True, size=10,
                              align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=False,
                              line_spacing=1.1)

    last = add_para_after(last,
        "Phần code trong các đoạn A.1 - A.3 trích lược từ tệp es_index.py và modules/doc_search.py của mã nguồn đề án "
        "(github.com/DatBa2/do-an-RAG). Trong báo cáo hoàn chỉnh sẽ trình bày kết quả _explain trên một câu truy vấn mẫu để "
        "đối chiếu trực tiếp với lý thuyết BM25 - RRF trình bày ở Chương 1.",
        italic=True)

    print("[+] Đã chèn Phụ lục A (mapping JSON + hybrid query)")

    doc.save(SRC)
    print(f"[OK] Saved {SRC}")


if __name__ == "__main__":
    main()
