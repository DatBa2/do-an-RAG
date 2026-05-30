"""Hậu xử lý file Decuong_NguyenBaDat_v2.docx:
1. Format toàn bộ font Times New Roman 13pt, justify, line-spacing 1.5, first-line indent 1cm
2. Căn giữa trang bìa
3. Chèn 3 ảnh sơ đồ vào đúng vị trí
4. Insert TOC field (mục lục tự động) - user nhấn F9 trong Word để cập nhật
"""
import copy
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = "/Users/badat/do_an_ths/do-an-RAG"
SRC = os.path.join(ROOT, "Decuong_NguyenBaDat_v2.docx")
DST = SRC  # ghi đè

ASSETS = os.path.join(ROOT, ".decuong_assets")
FIG1 = os.path.join(ASSETS, "fig_1_architecture.png")
FIG2 = os.path.join(ASSETS, "fig_2_sequence_docs.png")
FIG3 = os.path.join(ASSETS, "fig_3_sequence_function.png")

FONT = "Times New Roman"


def set_run_font(run, size_pt=13, bold=None, italic=None, color=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, indent_first=True, justify=True, line_spacing=1.5):
    pf = p.paragraph_format
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = line_spacing
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    if indent_first:
        pf.first_line_indent = Cm(1)


def is_centered_titlepage(para_text):
    """Heuristic: đoạn thuộc trang bìa (1, 2)."""
    keys = [
        "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH",
        "Nguyễn Bá Đạt",
        "NGHIÊN CỨU VÀ ỨNG DỤNG",
        "ĐỀ CƯƠNG ĐỀ ÁN TỐT NGHIỆP",
        "(Theo định hướng ứng dụng)",
        "HÀ NỘI",
        "CHUYÊN NGÀNH:",
        "MÃ SỐ:",
        "NGƯỜI HƯỚNG DẪN KHOA HỌC:",
        "---------------------------------------",
    ]
    txt = para_text.strip()
    return any(k in txt for k in keys)


def detect_titlepage_range(doc):
    """Đếm số paragraph thuộc 2 trang bìa đầu (trước heading MỤC LỤC)."""
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper().startswith("MỤC LỤC"):
            end_idx = i
            break
    return end_idx if end_idx is not None else 0


def format_all_paragraphs(doc):
    cover_end = detect_titlepage_range(doc)
    print(f"[i] Cover page ends at paragraph index {cover_end}")

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Phát hiện heading bằng style name
        style_name = (p.style.name or "").lower()
        is_heading = style_name.startswith("heading")
        level = 0
        if is_heading:
            for c in style_name:
                if c.isdigit():
                    level = int(c)
                    break

        # === Trang bìa ===
        if i < cover_end and (is_centered_titlepage(text) or text == ""):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.space_after = Pt(6)
            pf.line_spacing = 1.15
            for run in p.runs:
                # giữ size đã set khi tạo, chỉ thay font
                size = run.font.size or Pt(13)
                set_run_font(run, size_pt=size.pt if size else 13,
                             bold=run.bold, italic=run.italic)
            continue

        # === Heading ===
        if is_heading:
            for run in p.runs:
                if level == 1:
                    set_run_font(run, size_pt=14, bold=True)
                elif level == 2:
                    set_run_font(run, size_pt=13, bold=True)
                else:
                    set_run_font(run, size_pt=13, bold=True, italic=True)
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.3
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # === Body / list / table caption ===
        # Đoạn bắt đầu bằng '-' hay '(' hay 'FR' 'NFR' 'SEC' 'RQ' coi như bullet/list - không thụt
        no_indent = False
        for prefix in ("-", "FR", "NFR", "SEC", "RQ", "M1", "M2", "M3", "M4", "M5",
                       "UC0", "UC1", "(1)", "(2)", "(3)", "(4)", "(5)", "(a)", "(b)", "(c)",
                       "(M1", "(M2", "(M3", "(M4", "(M5", "[", "Mục tiêu",
                       "Yêu cầu chức năng", "Yêu cầu phi chức năng",
                       "Đối tượng nghiên cứu", "Phạm vi nghiên cứu",
                       "Tài liệu tiếng", "Giới thiệu chương", "Kết luận chương"):
            if text.startswith(prefix):
                no_indent = True
                break

        # Bỏ trống → không format
        if not text:
            continue

        for run in p.runs:
            set_run_font(run, size_pt=13, bold=run.bold, italic=run.italic)

        set_paragraph_format(p, indent_first=not no_indent,
                             justify=True, line_spacing=1.5)


def format_tables(doc):
    """Áp font TNR 12pt cho ô bảng, header bold, center."""
    for tbl in doc.tables:
        for r_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = None
                    pf.space_after = Pt(2)
                    pf.line_spacing = 1.15
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        set_run_font(run, size_pt=12,
                                     bold=(r_idx == 0) or run.bold,
                                     italic=run.italic)
                if r_idx == 0:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_paragraph_after(reference_para, text="", style=None):
    """Tạo paragraph mới ngay sau reference_para. Trả về paragraph mới."""
    new_p = OxmlElement("w:p")
    reference_para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, reference_para._parent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=13)
    if style:
        p.style = style
    return p


def add_picture_after(reference_para, image_path, width_cm=15, caption=""):
    """Chèn ảnh + caption ngay sau reference_para."""
    # 1. Tạo paragraph chứa ảnh
    p_pic = insert_paragraph_after(reference_para, text="")
    p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pic.paragraph_format.first_line_indent = None
    run = p_pic.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    # 2. Caption
    p_cap = insert_paragraph_after(p_pic, text=caption)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(12)
    for run in p_cap.runs:
        set_run_font(run, size_pt=12, italic=True, bold=True)
    return p_cap


def find_paragraph_by_text(doc, contains, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if contains in p.text:
            return i, p
    return -1, None


def insert_images(doc):
    """Chèn 3 ảnh vào đúng mục."""
    # Sơ đồ 1: sau đoạn '[Người dùng] → [Telegram /...]' (mục 2.6 kiến trúc)
    _, ref = find_paragraph_by_text(doc, "[Người dùng]")
    if ref is not None:
        add_picture_after(ref, FIG1, width_cm=15,
                          caption="Hình 2.1. Kiến trúc tổng thể 4 lớp của hệ thống.")
        print("[+] Inserted FIG1 (architecture) after Section 2.6")

    # Sơ đồ 2 + 3: cuối mục 3.4 (vòng lặp Advisor)
    # Chèn vào sau đoạn "Hệ chỉ thị (system instruction) được thiết kế..."
    _, ref = find_paragraph_by_text(doc, "Hệ chỉ thị (system instruction)")
    if ref is not None:
        last = add_picture_after(ref, FIG2, width_cm=15,
                                  caption="Hình 3.1. Sequence khi xử lý câu hỏi tài liệu nội bộ với Hybrid RAG (BM25 + KNN + RRF).")
        add_picture_after(last, FIG3, width_cm=15,
                          caption="Hình 3.2. Sequence khi xử lý câu hỏi dữ liệu cấu trúc với Function Calling và cổng RBAC.")
        print("[+] Inserted FIG2 + FIG3 (sequences) after Section 3.4")


def insert_toc_field(doc):
    """Thay placeholder DANH MỤC sau heading MỤC LỤC bằng TOC field thực."""
    # Tìm heading MỤC LỤC
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "MỤC LỤC":
            target_idx = i
            break
    if target_idx < 0:
        print("[!] Không tìm thấy heading MỤC LỤC")
        return

    # XÓA các paragraph placeholder kế tiếp tới khi gặp 'DANH MỤC CÁC KÝ HIỆU'
    paras = doc.paragraphs
    to_remove = []
    for j in range(target_idx + 1, len(paras)):
        if "DANH MỤC CÁC KÝ HIỆU" in paras[j].text:
            break
        to_remove.append(paras[j]._element)
    for el in to_remove:
        el.getparent().remove(el)
    print(f"[i] Xóa {len(to_remove)} placeholder paragraphs trong Mục lục")

    # Chèn TOC field ngay sau heading MỤC LỤC
    heading_para = doc.paragraphs[target_idx]
    new_p = OxmlElement("w:p")
    heading_para._element.addnext(new_p)

    # <w:r>...</w:r> chứa TOC field
    r1 = OxmlElement("w:r")
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    r1.append(fldChar1)
    new_p.append(r1)

    r2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r2.append(instrText)
    new_p.append(r2)

    r3 = OxmlElement("w:r")
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    r3.append(fldChar2)
    new_p.append(r3)

    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Mở file trong Microsoft Word, nhấn F9 hoặc bấm chuột phải vào đây và chọn 'Update Field' để hiển thị Mục lục."
    r4.append(t)
    new_p.append(r4)

    r5 = OxmlElement("w:r")
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    r5.append(fldChar3)
    new_p.append(r5)

    print("[+] Đã chèn TOC field. User cần nhấn F9 trong Word để cập nhật.")


def set_page_margins(doc, top=2.0, bottom=2.0, left=3.0, right=2.0):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def main():
    doc = Document(SRC)
    print(f"[i] Đọc {SRC} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")

    set_page_margins(doc)
    insert_images(doc)
    insert_toc_field(doc)
    format_all_paragraphs(doc)
    format_tables(doc)

    doc.save(DST)
    print(f"[OK] Đã ghi {DST}")


if __name__ == "__main__":
    main()
