"""Xóa các paragraph placeholder TOC giả nằm giữa TOC field thật và heading DANH MỤC CÁC KÝ HIỆU."""
from docx import Document
from docx.oxml.ns import qn

SRC = "/Users/badat/do_an_ths/do-an-RAG/Decuong_NguyenBaDat_v2.docx"

doc = Document(SRC)

# Tìm vị trí TOC field
toc_idx = -1
for i, p in enumerate(doc.paragraphs):
    for r in p._element.iter(qn("w:instrText")):
        if r.text and r.text.strip().startswith("TOC"):
            toc_idx = i
            break
    if toc_idx >= 0:
        break

print(f"TOC field at paragraph {toc_idx}")

# Tìm heading DANH MỤC CÁC KÝ HIỆU (style Heading 1) sau TOC
target_idx = -1
for j, p in enumerate(doc.paragraphs[toc_idx + 1:], start=toc_idx + 1):
    if (p.style.name or "").startswith("Heading") and "DANH MỤC CÁC KÝ HIỆU" in p.text:
        target_idx = j
        break

print(f"DANH MỤC heading at paragraph {target_idx}")

if toc_idx >= 0 and target_idx > toc_idx + 1:
    paras = doc.paragraphs
    to_remove = [paras[k]._element for k in range(toc_idx + 1, target_idx)]
    for el in to_remove:
        el.getparent().remove(el)
    print(f"Đã xóa {len(to_remove)} placeholder paragraphs giữa TOC và DANH MỤC")
else:
    print("Không cần xóa.")

doc.save(SRC)
print("OK")
