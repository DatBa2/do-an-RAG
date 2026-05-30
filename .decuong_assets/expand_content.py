"""Bổ sung lý thuyết, bảng và sơ đồ vào docx.

Chèn nội dung mở rộng vào các vị trí cụ thể, đồng thời chèn 9 sơ đồ bổ sung.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = "/Users/badat/do_an_ths/do-an-RAG"
SRC = os.path.join(ROOT, "Decuong_NguyenBaDat_v2.docx")
ASSETS = os.path.join(ROOT, ".decuong_assets")
FONT = "Times New Roman"


def set_run_font(run, size_pt=13, bold=None, italic=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if size_pt is not None: run.font.size = Pt(size_pt)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def add_para_after(ref, text="", bold=False, italic=False, size=13,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True, style=None):
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    p = Paragraph(new, ref._parent)
    if style:
        p.style = style
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size, bold=bold, italic=italic)
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if indent_first and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Cm(1)
    else:
        pf.first_line_indent = None
    return p


def add_picture_after(ref, image_path, caption, width_cm=15.5):
    p_pic = add_para_after(ref, "", align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    run = p_pic.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    p_cap = add_para_after(p_pic, caption, italic=True, bold=True, size=12,
                            align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    return p_cap


def add_table_after(ref, data, caption=None, header_bold=True):
    """Chèn bảng sau ref. data = list[list[str]]."""
    # Insert empty p first then table after that p
    p = add_para_after(ref, "", indent_first=False)
    rows = len(data); cols = len(data[0])
    tbl = ref._parent.add_table(rows=rows, cols=cols, width=Cm(16))
    tbl.style = "Table Grid"
    # Move table next to p
    p._element.addnext(tbl._element)
    p._element.getparent().remove(p._element)
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            cp = cell.paragraphs[0]
            run = cp.add_run(str(cell_text))
            set_run_font(run, size_pt=12, bold=(header_bold and i == 0))
            cp.paragraph_format.first_line_indent = None
            cp.paragraph_format.line_spacing = 1.15
            cp.paragraph_format.space_after = Pt(2)
    last = tbl
    if caption:
        # Insert caption paragraph after table
        new_p = OxmlElement("w:p")
        tbl._element.addnext(new_p)
        cp = Paragraph(new_p, ref._parent)
        run = cp.add_run(caption)
        set_run_font(run, size_pt=12, italic=True, bold=True)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = None
        cp.paragraph_format.space_after = Pt(12)
        return cp
    return tbl


def find_p(doc, contains, start=0, exact=False):
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        txt = p.text.strip()
        if (exact and txt == contains) or (not exact and contains in p.text):
            return i, p
    return -1, None


# =================================================================
# MAIN
# =================================================================
def main():
    doc = Document(SRC)

    # ---------------------------------------------------------------
    # 1) Thêm pipeline (FIG 4) sau mục 1.2 (kiến trúc RAG)
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Quy trình hoạt động chuẩn của một hệ RAG đầy đủ")
    if ref is not None:
        last = add_picture_after(ref, os.path.join(ASSETS, "fig_4_pipeline.png"),
                                  "Hình 1.1. Quy trình bốn pha của một hệ RAG (Advanced RAG).")
        # Thêm bảng so sánh 3 thế hệ RAG
        ref2 = last
        data = [
            ["Thế hệ", "Đặc trưng", "Ưu điểm", "Nhược điểm", "Ví dụ ứng dụng"],
            ["Naive RAG", "Truy xuất 1 lần + sinh 1 lần", "Đơn giản, tốc độ nhanh", "Nhạy cảm chất lượng retriever, ảo giác cao", "Chatbot Q&A tài liệu cơ bản"],
            ["Advanced RAG", "Pre/post-retrieval: rewrite, rerank, verify", "Chính xác cao, citation tốt", "Phức tạp hơn, độ trễ tăng", "Trợ lý nội bộ doanh nghiệp (đề án này)"],
            ["Modular RAG", "Mô-đun hoán đổi: Search/Memory/Routing/Fusion", "Linh hoạt, đa nguồn", "Yêu cầu kỹ năng kiến trúc cao", "Hệ thống đa-domain quy mô lớn"],
            ["Agentic RAG", "LLM tự lập kế hoạch + tool use", "Xử lý câu hỏi phức tạp", "Khó debug, chi phí token cao", "Trợ lý điều tra, multi-hop QA"],
        ]
        last = add_table_after(ref2, data,
            "Bảng 1.1. So sánh bốn thế hệ kiến trúc RAG (Gao et al., 2024; Singh et al., 2025).")
        # Đoạn lý thuyết mở rộng về 2 thành phần
        last = add_para_after(last,
            "Một góc nhìn bổ sung phân biệt hai trục thiết kế độc lập: trục retriever (sparse, dense, hybrid, graph-based) "
            "và trục generator (decoder-only LLM với hoặc không có function calling). Đề án này áp dụng hybrid retriever ở "
            "phía tài liệu, hybrid retriever ngầm thông qua function calling ở phía dữ liệu cấu trúc, và một generator duy "
            "nhất là Gemini 2.5 đảm nhận cả hai luồng. Cách tiếp cận này được Lewis (2020) gợi mở và Gao (2024) hệ thống "
            "hóa lại trong khung Modular RAG, nơi Search và Predict là hai mô-đun có thể hoán đổi.")
        print("[+] Mục 1.2: pipeline + bảng so sánh + đoạn mở rộng")

    # ---------------------------------------------------------------
    # 2) Mục 1.3 - thêm FIG 5 và FIG 6, đoạn lý thuyết DPR/ColBERT
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Với tiếng Việt, lựa chọn mô hình embedding")
    if ref is not None:
        last = add_para_after(ref,
            "Trong dòng nghiên cứu dense retrieval, Karpukhin và cộng sự (2020) đặt nền với mô hình DPR sử dụng hai encoder "
            "tách biệt cho câu hỏi và đoạn văn cùng huấn luyện đối chiếu trên cặp (Q, P+). Khanh và cộng sự (2020) đề xuất "
            "ColBERT với late interaction: thay vì biểu diễn cả tài liệu bằng một vector, mỗi token sinh một vector và truy "
            "xuất so khớp tối đa qua MaxSim - cải thiện đáng kể độ chính xác với chi phí lưu trữ lớn hơn. Cho tiếng Việt, "
            "VinAI phát hành PhoBERT (Nguyen, 2020) - mô hình monolingual đầu tiên đạt SOTA trên các tác vụ POS tagging, NER, "
            "NLI; biến thể sup-SimCSE-Vietnamese-PhoBERT-base sau đó được tinh chỉnh cho retrieval. Benchmark VN-MTEB (Tran, "
            "2025) công bố tháng 7/2025 cho thấy các mô hình đa ngôn ngữ thế hệ mới (text-embedding-004, bge-m3, jina-v3) "
            "đã thu hẹp khoảng cách với mô hình bản địa nhờ corpus huấn luyện rộng và kỹ thuật contrastive trên đa ngôn ngữ.")
    _, ref = find_p(doc, "Hạn chế nằm ở chỗ BM25 không hiểu ngữ nghĩa")
    if ref is not None:
        add_picture_after(ref, os.path.join(ASSETS, "fig_9_inverted_index.png"),
                          "Hình 1.2. Minh hoạ chỉ mục đảo (Inverted Index) và cách BM25 vận hành trên Elasticsearch.")
    _, ref = find_p(doc, "Reciprocal Rank Fusion (Cormack, Clarke và Buettcher")
    if ref is not None:
        last = add_picture_after(ref, os.path.join(ASSETS, "fig_5_retrieval_modes.png"),
                                  "Hình 1.3. So sánh ba phương pháp truy xuất sparse, dense và hybrid.")
        last = add_picture_after(last, os.path.join(ASSETS, "fig_6_rrf_example.png"),
                                  "Hình 1.4. Ví dụ minh hoạ thuật toán Reciprocal Rank Fusion với k = 60.")
        # Bảng so sánh embedding tiếng Việt
        data = [
            ["Mô hình embedding", "Loại", "Chiều", "Nguồn / năm", "Ghi chú"],
            ["text-embedding-004", "Đa ngôn ngữ", "768", "Google, 2024", "Hai mode retrieval_document / _query"],
            ["bge-m3", "Đa ngôn ngữ", "1024", "BAAI, 2024", "Hỗ trợ multi-vector + sparse"],
            ["jina-embeddings-v3", "Đa ngôn ngữ", "1024", "Jina AI, 2024", "Task-aware adapter"],
            ["PhoBERT (base/large)", "Bản địa VN", "768/1024", "VinAI, 2020", "Cần word segmentation"],
            ["sup-SimCSE-Vietnamese-PhoBERT-base", "Bản địa VN", "768", "Cộng đồng, 2022", "Fine-tune cho retrieval"],
            ["Vietnamese-Embedding (vinai/...)", "Bản địa VN", "768", "VinAI, 2023", "Tối ưu Vietnamese retrieval"],
        ]
        last = add_table_after(last, data,
            "Bảng 1.2. So sánh các mô hình embedding khả dụng cho tiếng Việt (cập nhật theo VN-MTEB 2025).")
        # Đoạn về rerank
        last = add_para_after(last,
            "Sau bước RRF, một số hệ thống bổ sung tầng xếp hạng lại (rerank) bằng cross-encoder - mô hình nhận đồng thời "
            "(câu hỏi, đoạn) và sinh điểm tương đồng. Cross-encoder cho điểm chính xác hơn bi-encoder nhưng chi phí cao do "
            "phải chạy với mỗi cặp. Trong đề án này, rerank không được mặc định bật mà được liệt vào trục ablation: bật/tắt "
            "để đo chi phí - lợi ích trên các nhóm câu hỏi dài và phức tạp.")
        print("[+] Mục 1.3: thêm 3 ảnh + bảng embedding + đoạn DPR/ColBERT/rerank")

    # ---------------------------------------------------------------
    # 3) Mục 1.4 - thêm FIG 10 HNSW + bảng vector DB + đoạn về quantization
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Đối với trường dense_vector, mapping khai báo chiều 768")
    if ref is not None:
        last = add_picture_after(ref, os.path.join(ASSETS, "fig_10_hnsw.png"),
                                  "Hình 1.5. Cấu trúc HNSW: đồ thị phân tầng giúp tìm KNN xấp xỉ với độ phức tạp dưới tuyến tính.")
        last = add_para_after(last,
            "Cơ sở lý thuyết của HNSW (Malkov & Yashunin, 2018) là kết hợp hai ý tưởng: đồ thị 'small-world' của Kleinberg "
            "(ngẫu nhiên hoá liên kết để rút ngắn quãng đường) và cấu trúc skip-list phân tầng. Mỗi tầng cao chứa thưa hơn, "
            "các tầng thấp dày đặc. Truy vấn bắt đầu ở tầng cao nhất, đi greedy đến láng giềng gần nhất trong tầng đó, rồi "
            "tụt xuống tầng thấp hơn lặp lại cho đến khi tới tầng 0 - nơi chứa toàn bộ vector. Hai siêu tham số quan trọng: "
            "m kiểm soát độ kết nối (mặc định 16) và ef_construction điều khiển chất lượng build (100). Phía truy vấn, "
            "ef_search điều chỉnh đánh đổi độ chính xác - tốc độ. Từ Elasticsearch 8.12, kiểu vector int8_hnsw cho phép nén "
            "vector từ float32 xuống int8, giảm 4 lần dung lượng đĩa với mất mát chính xác thường dưới 1 điểm phần trăm trên "
            "benchmark MTEB.")
        # Bảng so sánh vector DB
        data = [
            ["Hệ thống", "Loại lưu trữ", "Hỗ trợ ANN", "BM25 native", "Hybrid native", "Ghi chú"],
            ["Elasticsearch 8.x", "Document store + ANN", "HNSW (int8/float)", "Có", "RRF qua retriever API", "Đề án dùng"],
            ["OpenSearch", "Document store + ANN", "HNSW, IVF", "Có", "RRF từ 2.11", "Fork ES"],
            ["Qdrant", "Vector DB chuyên dụng", "HNSW", "Không", "Có (BM42 mới)", "Rust, fast"],
            ["Milvus / Zilliz", "Vector DB chuyên dụng", "HNSW, IVF, DiskANN", "Không", "Sparse + dense", "Quy mô tỷ"],
            ["Pinecone", "Managed cloud", "Proprietary", "Không", "Sparse-dense hybrid", "Pay-as-you-go"],
            ["FAISS", "Library", "HNSW, IVF, PQ", "Không", "Không", "Local, Meta"],
            ["pgvector / Tiger", "PostgreSQL extension", "HNSW, IVFFlat", "Có (tsvector)", "Đang phát triển", "Tích hợp SQL"],
        ]
        last = add_table_after(last, data,
            "Bảng 1.3. So sánh các giải pháp lưu trữ vector phổ biến 2024-2025.")
        print("[+] Mục 1.4: HNSW + bảng vector DB + đoạn quantization")

    # ---------------------------------------------------------------
    # 4) Mục 1.5 - thêm đoạn về Transformer + prompting + tools
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Function calling là cơ chế cho phép LLM")
    if ref is not None:
        last = add_para_after(ref,
            "Về mặt kỹ thuật, function calling được hiện thực thông qua schema JSON chuẩn (OpenAPI-like): mô hình được cung "
            "cấp danh sách FunctionDeclaration kèm mô tả ngôn ngữ tự nhiên và schema tham số (kiểu, bắt buộc/tuỳ chọn, "
            "enum, ràng buộc). Khi nhận câu hỏi, LLM có ba lựa chọn: trả văn bản trực tiếp, gọi một hàm với tham số JSON, "
            "hoặc kết hợp - nhiều lượt gọi tuần tự (ReAct). Trong đề án này, vòng lặp dispatcher giới hạn MAX_TOOL_LOOPS=8 "
            "để chống đệ quy vô hạn khi LLM gọi tool sai liên tục. Đây cũng là khung mở rộng tự nhiên: thêm tool mới chỉ "
            "cần đăng ký FunctionDeclaration và mapping Python tương ứng.")
    # Thêm bảng các kỹ thuật prompting + đoạn về safety
    _, ref = find_p(doc, "Đề án lựa chọn Gemini 2.5 Flash làm mô hình mặc định")
    if ref is not None:
        last = add_para_after(ref,
            "Bên cạnh function calling, các kỹ thuật prompting được vận dụng trong hệ chỉ thị (system instruction) gồm: "
            "instruction prompting (nêu rõ mục tiêu, ngôn ngữ trả lời, định dạng), few-shot prompting (đưa 1-2 ví dụ mẫu cho "
            "câu hỏi mơ hồ), RAG-aware prompting (nhấn 'chỉ dùng ngữ cảnh được cung cấp'), và bộ khung từ chối an toàn "
            "(nếu câu hỏi vượt phạm vi nội bộ, mô hình từ chối lịch sự chứ không suy diễn).")
        data = [
            ["Kỹ thuật prompting", "Ý tưởng cốt lõi", "Áp dụng trong đề án"],
            ["Zero-shot", "Đặt câu hỏi không ví dụ", "Mặc định cho mọi câu hỏi"],
            ["Few-shot", "Cung cấp 1-N cặp (Q, A) mẫu", "Trong prompt verifier và rewrite"],
            ["Chain-of-Thought", "Yêu cầu lập luận từng bước", "Không bật cho user-facing answer (giữ ngắn gọn)"],
            ["Self-consistency", "Lấy nhiều mẫu, chọn đa số", "Không dùng (tốn token)"],
            ["RAG-aware", "Bắt buộc bám ngữ cảnh, trích nguồn", "Là khung system instruction"],
            ["Function calling", "LLM phát lời gọi hàm JSON", "Cốt lõi cho 20 tool"],
            ["Self-RAG / verify", "Mô hình tự kiểm tra grounding", "Phiên bản lite ở Advisor._verify()"],
        ]
        last = add_table_after(last, data,
            "Bảng 1.4. Các kỹ thuật prompting phổ biến và mức độ áp dụng trong đề án.")
        print("[+] Mục 1.5: thêm 2 đoạn + bảng prompting")

    # ---------------------------------------------------------------
    # 5) Mục 1.6 - thêm FIG 7 chunking + bảng chiến lược + đoạn về MetaRAG
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Khuyến nghị công nghiệp 2025 là chia chunk khoảng 256-512")
    if ref is not None:
        last = add_picture_after(ref, os.path.join(ASSETS, "fig_7_chunking.png"),
                                  "Hình 1.6. Bốn chiến lược chunking phổ biến và đánh giá tóm tắt.")
        data = [
            ["Chiến lược", "Cơ chế", "Ưu", "Nhược", "Phù hợp"],
            ["Fixed-size", "Cắt theo độ dài cố định", "Nhanh, đơn giản, ít tham số", "Đứt câu, mất ngữ cảnh", "Tài liệu thuần văn"],
            ["Recursive", "Đệ quy theo dấu phân cấp", "Giữ cấu trúc tự nhiên", "Khó với tài liệu không heading", "Báo cáo, sách giáo trình"],
            ["Section-aware", "Cắt theo heading + overlap", "Mỗi chunk ≈ 1 chủ đề", "Phụ thuộc heading rõ", "Markdown, Wiki (đề án dùng)"],
            ["Semantic", "Gom theo embedding similarity", "Coherence cao", "Chi phí gấp đôi, không vượt fixed (NAACL 2025)", "Tài liệu không cấu trúc"],
            ["Hierarchical", "Đa tầng: tóm tắt + chi tiết", "Tốt cho multi-hop", "Phức tạp, tốn lưu trữ", "Kho tài liệu lớn"],
            ["Late chunking", "Embed cả tài liệu rồi cắt vector", "Giữ ngữ cảnh dài", "Yêu cầu mô hình context dài", "Tài liệu dày"],
            ["Agentic", "LLM tự quyết chỗ cắt", "Linh hoạt nhất", "Đắt nhất", "Thử nghiệm"],
        ]
        last = add_table_after(last, data,
            "Bảng 1.5. Các chiến lược chunking và đánh giá so sánh.")
        last = add_para_after(last,
            "Về tự xác thực (Self-RAG), các nghiên cứu 2025 đưa ra nhiều hướng đi không cần fine-tune mô hình. MetaRAG (Vol-"
            "4136 IAAI 2025) áp dụng metamorphic testing: sinh các biến thể câu trả lời bằng phép thay đồng nghĩa/trái nghĩa, "
            "đối chiếu với ngữ cảnh - phát hiện ảo giác ở chế độ black-box thời gian thực. HalluGraph (arXiv:2512.01659) đề "
            "xuất khung đồ thị tách câu trả lời thành các bộ ba (entity, relation, entity) và đo độ trùng với knowledge graph "
            "ngữ cảnh. RT4CHART (arXiv:2603.27752) chia kiểm tra thành hai pha local và global trên từng atomic claim. Đề án "
            "lựa chọn cách triển khai đơn giản hơn (verify-by-prompt một lần) để cân bằng giữa độ tin cậy và chi phí token.")
        print("[+] Mục 1.6: ảnh chunking + bảng + đoạn MetaRAG/HalluGraph")

    # ---------------------------------------------------------------
    # 6) Mục 1.7 - thêm FIG 8 RBAC + bảng OWASP Top 10 đầy đủ
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Kiểm soát truy cập theo vai trò (RBAC) là kĩ thuật nền")
    if ref is not None:
        last = add_picture_after(ref, os.path.join(ASSETS, "fig_8_rbac.png"),
                                  "Hình 1.7. Cơ chế phân quyền RBAC hai lớp (pre-filter + post-gate) phản ánh khuyến nghị OWASP LLM02 - 2025.")
        data = [
            ["Mã", "Tên", "Mô tả", "Áp dụng đối với đề án"],
            ["LLM01", "Prompt Injection", "Kẻ tấn công chèn chỉ thị độc trực/gián tiếp", "System prompt tách dữ liệu; lọc câu hỏi"],
            ["LLM02", "Sensitive Information Disclosure", "Rò rỉ PII, secrets qua câu trả lời", "ACL allowed_roles + post-filter"],
            ["LLM03", "Supply Chain", "Vuln từ model, dataset, library bên thứ ba", "Pin version, audit dependency"],
            ["LLM04", "Data and Model Poisoning", "Đầu độc dữ liệu huấn luyện hoặc kho RAG", "Whitelist nguồn tài liệu admin"],
            ["LLM05", "Improper Output Handling", "Không sanitize output → XSS/SSRF/code injection", "Render HTML có escape; không exec"],
            ["LLM06", "Excessive Agency", "Tác tử có quá nhiều quyền hành động", "Giới hạn 20 tool đọc, không ghi/xoá"],
            ["LLM07", "System Prompt Leakage", "Lộ system prompt nhạy cảm", "Không nhúng secrets trong prompt"],
            ["LLM08", "Vector and Embedding Weaknesses", "Mới 2025: tấn công vector store", "Lưu allowed_roles ở cùng vector"],
            ["LLM09", "Misinformation", "Trả lời sai do nguồn kém", "Self-RAG verify + citation bắt buộc"],
            ["LLM10", "Unbounded Consumption", "Token, request bị khai thác trục lợi", "Rate-limit Telegram + audit log"],
        ]
        last = add_table_after(last, data,
            "Bảng 1.6. OWASP Top 10 cho LLM Applications phiên bản 2025 và mức ánh xạ với đề án.")
        last = add_para_after(last,
            "Một xu hướng mới năm 2025 là sử dụng chính LLM làm bộ phân loại quyền truy cập trong RAG, được khung ARBITER "
            "(arXiv:2512.20535) đề xuất với NeMo Guardrails: thay vì duy trì bảng ACL tĩnh, một LLM nhẹ phán đoán quyền hiển "
            "thị nội dung theo few-shot prompt. Đây là hướng linh hoạt nhưng kém deterministic; đề án giữ ACL tĩnh "
            "(allowed_roles trong mapping ES) để có thể kiểm chứng được - phù hợp ngữ cảnh giáo dục công nơi mọi quyết định "
            "lộ thông tin cần truy vết được.")
        print("[+] Mục 1.7: FIG 8 RBAC + bảng OWASP đầy đủ + ARBITER")

    # ---------------------------------------------------------------
    # 7) Mục 1.8 - bảng RAGAS metrics chi tiết
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "Ngoài RAGAS, các nghiên cứu IR truyền thống dùng")
    if ref is not None:
        data = [
            ["Chỉ số", "Nhóm", "Định nghĩa rút gọn", "Khoảng giá trị", "Tham chiếu"],
            ["Faithfulness", "Generation", "% tuyên bố trong answer được context hậu thuẫn", "[0, 1] - cao tốt", "Es et al. 2024"],
            ["Answer Relevancy", "Generation", "Cosine giữa câu hỏi gốc và câu hỏi tái sinh từ answer", "[0, 1] - cao tốt", "Es et al. 2024"],
            ["Context Precision", "Retrieval", "% chunk được rank cao đúng là liên quan", "[0, 1] - cao tốt", "Es et al. 2024"],
            ["Context Recall", "Retrieval", "% câu trong ground-truth có chunk hỗ trợ", "[0, 1] - cao tốt", "Es et al. 2024"],
            ["Context Entity Recall", "Retrieval", "% entity trong ground-truth xuất hiện trong context", "[0, 1] - cao tốt", "Es et al. 2024"],
            ["Answer Correctness", "End-to-end", "Tương đương ngữ nghĩa + factual với reference", "[0, 1] - cao tốt", "Es et al. 2024"],
            ["Recall@k / MRR / nDCG@k", "Retrieval cổ điển", "Bao phủ / hạng trung bình / discounted gain", "[0, 1]", "TREC"],
            ["F1, BLEU, ROUGE, BERTScore", "Generation cổ điển", "So khớp n-gram hoặc embedding với reference", "[0, 1]", "NLP chuẩn"],
            ["ACL compliance (đề án)", "Bảo mật", "expected_denied_source KHÔNG xuất hiện trong citations", "0/1", "Đề án này"],
            ["Denial handle (đề án)", "Bảo mật", "Tỉ lệ trả thông báo từ chối đúng khi vượt quyền", "[0, 1]", "Đề án này"],
        ]
        last = add_table_after(ref, data,
            "Bảng 1.7. Bộ chỉ số RAGAS, IR và bộ chỉ số bổ sung do đề án định nghĩa.")
        last = add_para_after(last,
            "Ngoài các chỉ số chính xác, các tổ chức triển khai sản phẩm chú trọng thêm chỉ số vận hành: độ trễ phân vị "
            "p50/p95/p99, thông lượng (qps), số token in/out trung bình, tỉ lệ cache hit ở tầng embedding, tỉ lệ tool error. "
            "Đề án tính p50/p95 cho mỗi baseline và đo riêng chi phí trung bình của bước verify so với không verify - phục vụ "
            "kết luận về đánh đổi chính xác - chi phí.")
        print("[+] Mục 1.8: bảng RAGAS + đoạn về vận hành")

    # ---------------------------------------------------------------
    # 8) Mục 2.3 - thêm bảng synonym mapping
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "synonym_graph token filter")
    if ref is None:
        _, ref = find_p(doc, "Đặc trưng tiếng Việt cần xử lý gồm")
    if ref is not None:
        data = [
            ["Từ chuẩn", "Biến thể / viết tắt được map", "Bối cảnh"],
            ["học sinh", "HS, em, sinh viên (thận trọng)", "Toàn bộ nội dung"],
            ["học kỳ 1", "HK1, học kì 1, kỳ 1, hk1", "Điểm, thống kê"],
            ["học kỳ 2", "HK2, học kì 2, kỳ 2, hk2", "Điểm, thống kê"],
            ["trung bình", "TB, ĐTB, GPA, điểm trung bình, điểm tb", "Báo cáo điểm"],
            ["giáo viên chủ nhiệm", "GVCN, gvcn, chủ nhiệm", "Nhận xét"],
            ["thường xuyên", "TX, kiểm tra TX, hệ số 1", "Cột điểm"],
            ["giữa kỳ", "GK, kiểm tra GK, hệ số 2", "Cột điểm"],
            ["cuối kỳ", "CK, kiểm tra CK, hệ số 3", "Cột điểm"],
            ["tổng kết", "TK, điểm TK, điểm tổng", "Cột điểm"],
            ["hạnh kiểm", "HK (lưu ý ngữ cảnh), HK học sinh", "Đánh giá phẩm chất"],
            ["xếp loại", "loại, học lực, đánh giá", "Báo cáo cuối kỳ"],
            ["Thông tư 22", "TT22, TT 22/2021", "Văn bản pháp lý"],
        ]
        last = add_table_after(ref, data,
            "Bảng 2.1. Trích bảng synonym mapping tiếng Việt cho lĩnh vực giáo dục (cấu hình trong analyzer vn_text).")
        print("[+] Mục 2.3: bảng synonym mapping")

    # ---------------------------------------------------------------
    # 9) Mục 4.1-4.5 - thêm bảng chi tiết hơn + biểu đồ baseline + heatmap chunking
    # ---------------------------------------------------------------
    _, ref = find_p(doc, "RQ9 - Hệ thống có phục vụ được use case quản trị")
    if ref is not None:
        # Bảng mapping RQ → metric chính
        data = [
            ["RQ", "Mục tiêu kiểm chứng", "Chỉ số chính", "Nhóm test set"],
            ["RQ1", "Hybrid > từng phương pháp đơn", "Citation match, Context Precision", "B (docs)"],
            ["RQ2", "Function calling > vector cho cấu trúc", "Tool recall, Keyword recall", "A (structured)"],
            ["RQ3", "Query rewrite hỗ trợ multi-turn", "Tool recall trên G", "G (multi-turn)"],
            ["RQ4", "Self-verify giảm ảo giác", "Not-found handle, Hallucination rate", "J (hallucination)"],
            ["RQ5", "Tham số chunking ảnh hưởng kết quả", "Heatmap F1 trên B", "B + ablation chunking"],
            ["RQ6", "Robust với biến thể tiếng Việt", "Keyword recall trên D", "D (robustness)"],
            ["RQ7", "Chịu prompt injection", "Denial handle trên I, H", "I + H"],
            ["RQ8", "RBAC chặn vượt quyền", "ACL compliance", "L (rbac)"],
            ["RQ9", "Use case quản trị hoạt động", "Tool recall + Keyword recall", "M (admin)"],
        ]
        last = add_table_after(ref, data,
            "Bảng 4.1. Mapping câu hỏi nghiên cứu - chỉ số chính - nhóm test set.")
        print("[+] Mục 4.1: bảng RQ ↔ metric ↔ nhóm")

    # Bảng giả định kết quả baseline + chart
    _, ref = find_p(doc, "Hai trục liên quan tới chunking phải re-ingest tài liệu")
    if ref is not None:
        last = add_picture_after(ref, os.path.join(ASSETS, "fig_11_baseline_chart.png"),
                                  "Hình 4.1. Dự kiến phân bố điểm 4 chỉ số chính giữa 5 baseline (giá trị minh hoạ).")
        # Bảng số chi tiết baseline kỳ vọng
        data = [
            ["Baseline", "Tool recall", "Citation match", "Not-found handle", "ACL compliance", "Latency p50 (ms)", "Latency p95 (ms)"],
            ["raw_llm", "0.10", "0.05", "0.20", "0.30", "1500", "2600"],
            ["bm25_only", "0.55", "0.62", "0.55", "0.65", "2200", "3800"],
            ["vector_only", "0.55", "0.58", "0.50", "0.60", "2800", "4500"],
            ["function_only", "0.88", "0.70", "0.78", "0.92", "3100", "5200"],
            ["hybrid_full", "0.92", "0.86", "0.90", "1.00", "3500", "6500"],
        ]
        last = add_table_after(last, data,
            "Bảng 4.2. Bảng số dự kiến cho 5 baseline (giá trị minh hoạ, sẽ thay bằng kết quả thực).")
        # Heatmap chunking
        last = add_picture_after(last, os.path.join(ASSETS, "fig_12_chunking_heatmap.png"),
                                  "Hình 4.2. Dự kiến heatmap F1 theo (chunk_size × overlap) - sweet spot 320 × 64.")
        # Bảng risk vs mitigation
        data = [
            ["Rủi ro", "Mô tả", "Biện pháp giảm thiểu trong đề án"],
            ["Rò rỉ tài liệu admin/teacher cho parent", "Câu hỏi vô tình match tài liệu cấp cao", "Pre-filter allowed_roles ở ES + post-render filter"],
            ["Prompt injection trực tiếp", "User ép LLM bỏ system prompt", "System prompt khoá vai dữ liệu + lọc pattern"],
            ["Prompt injection gián tiếp qua tài liệu", "Tài liệu chứa chỉ thị độc", "Admin duyệt tài liệu trước khi ingest"],
            ["Token blow-up", "Câu hỏi rất dài làm context tăng", "Cắt theo MAX_CONTEXT, từ chối quá ngưỡng"],
            ["Tool loop vô hạn", "LLM lặp gọi tool sai", "MAX_TOOL_LOOPS=8 + fail-fast"],
            ["Embedding cache lỗi", "Trả vector cũ sau khi mô hình thay đổi", "Cache key gồm tên model + task_type"],
            ["Cluster ES quá tải", "Quá nhiều câu hỏi đồng thời", "Rate-limit Telegram + audit cảnh báo"],
            ["Sai chính tả PII", "Tên không dấu khác bản chính", "asciifolding + synonym + ambiguity handle"],
        ]
        last = add_table_after(last, data,
            "Bảng 4.3. Các rủi ro vận hành và biện pháp giảm thiểu đã có sẵn trong thiết kế.")
        print("[+] Mục 4.5: bảng baseline + heatmap + bảng risk")

    doc.save(SRC)
    print(f"[OK] Saved {SRC}")


if __name__ == "__main__":
    main()
